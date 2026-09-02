# -*- coding: utf-8 -*-
"""docx 파서 + stage0 디스패처 회귀 테스트.

    pytest tests/test_docx_extract.py -q

■ 무엇을 지키나
  ① 문단·표가 **문서에 나온 순서 그대로** 나온다 (표/문단을 따로 읽으면 순서가 깨진다)
  ② 표 셀 경계는 줄바꿈, 가로 병합 셀은 행 안에서 중복 제거된다
  ③ 🔴 확장자 위장 — `.docx` 인데 내용이 XLSX/PPTX/HWPX/ODF/정체불명이면
     조용히 빈 문자열이 아니라 정체를 담아 실패한다
  ④ `stage0_extract` 는 확장자가 아니라 **내용물**로 hwpx/docx 를 가른다 —
     `.hwpx` 확장자에 docx 내용이, `.docx` 확장자에 hwpx 내용이 와도 맞게 파싱된다

DB 를 쓰지 않는다(다른 세션들이 같은 Postgres 를 공유 중). 픽스처는 모두
`python-docx` 로 그 자리에서 만든 진짜 OOXML 이거나(모의가 아니다) 저장소에
이미 있는 실물 hwpx 다.
"""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx_extract as dx           # noqa: E402
import hwpx_extract as hx           # noqa: E402
import stage0_extract as s0         # noqa: E402
from docx import Document           # noqa: E402
from docx.enum.text import WD_BREAK  # noqa: E402

실물_HWPX = ROOT / "2026_Finance_DATA_FOR_RAG/창진원/초격차 스타트업 프로젝트/초격차 스타트업 프로젝트 세부관리기준(제10차).hwpx"


# ── ①②경계·순서 규칙 (python-docx 로 만든 진짜 docx) ─────────────────
def test_문단과_표가_문서_순서대로_나온다(tmp_path):
    d = Document()
    d.add_paragraph("앞 문단")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "총 사업비"
    t.cell(0, 1).text = "정부지원사업비"
    t.cell(1, 0).text = "100%"
    t.cell(1, 1).text = "70%"
    d.add_paragraph("뒤 문단")
    f = tmp_path / "합성.docx"
    d.save(f)

    assert dx.extract(f).split("\n") == [
        "앞 문단", "총 사업비", "정부지원사업비", "100%", "70%", "뒤 문단",
    ]


def test_lineBreak와_tab이_줄바꿈_탭으로_산다(tmp_path):
    d = Document()
    p = d.add_paragraph("줄1")
    p.add_run().add_break(WD_BREAK.LINE)
    p.add_run("줄2")
    p2 = d.add_paragraph("탭전")
    p2.add_run().add_tab()
    p2.add_run("탭후")
    f = tmp_path / "줄바꿈.docx"
    d.save(f)
    assert dx.extract(f).split("\n") == ["줄1", "줄2", "탭전\t탭후"]


def test_빈문단은_줄을_만들지_않는다(tmp_path):
    d = Document()
    d.add_paragraph("A")
    d.add_paragraph("")
    d.add_paragraph("   ")
    d.add_paragraph("B")
    f = tmp_path / "빈문단.docx"
    d.save(f)
    assert dx.extract(f) == "A\nB"


def test_가로병합_셀은_행_안에서_중복제거된다(tmp_path):
    d = Document()
    t = d.add_table(rows=2, cols=3)
    merged = t.cell(0, 0).merge(t.cell(0, 1))
    merged.text = "병합칸"
    t.cell(0, 2).text = "우측"
    t.cell(1, 0).text = "1행0"
    t.cell(1, 1).text = "1행1"
    t.cell(1, 2).text = "1행2"
    f = tmp_path / "병합.docx"
    d.save(f)
    assert dx.extract(f).split("\n") == ["병합칸", "우측", "1행0", "1행1", "1행2"]


def test_중첩_표도_셀_안에서_순서대로_나온다(tmp_path):
    d = Document()
    t = d.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    cell.paragraphs[0].text = "바깥칸 문단"
    inner_tbl = cell.add_table(rows=1, cols=1)
    inner_tbl.cell(0, 0).text = "안쪽칸"
    f = tmp_path / "중첩표.docx"
    d.save(f)
    text = dx.extract(f)
    assert "바깥칸 문단" in text and "안쪽칸" in text


# ── ③ 확장자 위장 방어 ──────────────────────────────────────────────
def _fake_zip(tmp_path: Path, name: str, entries: dict[str, str]) -> Path:
    f = tmp_path / name
    with zipfile.ZipFile(f, "w") as z:
        for k, v in entries.items():
            z.writestr(k, v)
    return f


@pytest.mark.parametrize("entries, 정체", [
    ({"[Content_Types].xml": "<x/>", "xl/workbook.xml": "<x/>"}, "XLSX"),
    ({"[Content_Types].xml": "<x/>", "ppt/presentation.xml": "<x/>"}, "PPTX"),
    ({"mimetype": hx.HWPX_MIMETYPE, "Contents/section0.xml": "<x/>"}, "hwpx"),
    ({"mimetype": "application/vnd.oasis.opendocument.text", "content.xml": "<x/>"}, "ODF"),
    ({"README.txt": "hi"}, "docx 아님"),
])
def test_위장_zip은_정체를_밝히며_실패한다(tmp_path, entries, 정체):
    f = _fake_zip(tmp_path, "위장.docx", entries)
    with pytest.raises(dx.NotDocxError) as ei:
        dx.extract(f)
    assert 정체 in str(ei.value) and "위장.docx" in str(ei.value)


def test_zip_아닌_docx는_매직바이트로_정체를_밝힌다(tmp_path):
    for head, 정체 in ((b"\xd0\xcf\x11\xe0" + b"\0" * 32, "OLE2"),
                       (b"%PDF-1.4\n", "PDF"),
                       (b"<?xml version='1.0'?><x/>", "XML"),
                       (b"\x00\x01garbage", "정체 불명")):
        f = tmp_path / "가짜.docx"
        f.write_bytes(head)
        with pytest.raises(dx.NotDocxError, match=정체):
            dx.extract(f)


# ── ④ stage0 디스패처 — 확장자가 아니라 내용물 ─────────────────────────
def test_stage0는_docx를_정상_파싱한다(tmp_path):
    d = Document()
    d.add_paragraph("제1조(목적) 시험용이다.")
    f = tmp_path / "규정.docx"
    d.save(f)
    kind, (text, offs) = s0.extract(f)
    assert kind == "text" and text == "제1조(목적) 시험용이다." and offs == {}


def test_stage0는_hwpx_확장자여도_내용이_docx면_docx로_판다(tmp_path):
    d = Document()
    d.add_paragraph("확장자는 hwpx 지만 진짜는 docx")
    f = tmp_path / "위장.hwpx"
    d.save(f)
    kind, (text, _) = s0.extract(f)
    assert text == "확장자는 hwpx 지만 진짜는 docx"


@pytest.mark.skipif(not 실물_HWPX.exists(), reason="실물 hwpx 픽스처 없음")
def test_stage0는_docx_확장자여도_내용이_hwpx면_hwpx로_판다(tmp_path):
    f = tmp_path / "위장.docx"
    f.write_bytes(실물_HWPX.read_bytes())
    kind, (text, _) = s0.extract(f)
    assert "제1조(목적)" in text


def test_stage0는_docx_확장자에_XLSX가_오면_정체를_밝히며_실패한다(tmp_path):
    f = _fake_zip(tmp_path, "위장.docx", {"xl/workbook.xml": "<x/>"})
    with pytest.raises(ValueError, match="XLSX"):
        s0.extract(f)


def test_doc_확장자는_지원하지_않는_형식으로_실패한다(tmp_path):
    """구형 .doc(OLE2) 파서는 만들지 않는다 — 제품 결정: 업로드 허용 목록에서 뺀다."""
    f = tmp_path / "구형.doc"
    f.write_bytes(b"\xd0\xcf\x11\xe0" + b"\0" * 32)
    with pytest.raises(ValueError, match="지원하지 않는 형식"):
        s0.extract(f)
