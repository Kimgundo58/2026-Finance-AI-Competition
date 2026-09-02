# -*- coding: utf-8 -*-
"""DOCX(OOXML WordprocessingML) 본문 텍스트 추출기 (python-docx 1.2.0).

시그니처·반환형은 `hwp_extract.extract(path) -> str` · `hwpx_extract.extract(path) -> str`
과 같다. 정체 판정은 새로 만들지 않고 `hwpx_extract.sniff()` 를 그대로 쓴다 —
docx 도 zip 이라 판정 로직이 같고, 두 벌로 갈라 쓰면 한쪽만 고쳐지는 날이 온다.

🔴 경계 규칙(hwpx 와 동일) — 문단 하나 = 한 줄, 표 셀 하나 = 한 줄(셀 안 문단이
   여럿이면 각각 한 줄). `python-docx` 의 `Paragraph.text` 가 `w:br`/`w:cr` 를 이미
   `\\n` 으로, `w:tab` 을 `\\t` 로 넣어 준다 — 그대로 살린다.

🔴 문서 순서 — 본문 문단과 표는 `Document.iter_inner_content()` 로 순회한다.
   `document.paragraphs`/`document.tables` 를 따로 읽으면 둘의 등장 순서가
   섞여 조문 순서가 깨진다.

🔴 확장자 위장 방어 — `.docx` 여도 zip 을 열어 `word/document.xml` 존재로 재검증한다.
   아니면(XLSX·PPTX·ODF·HWPX·정체불명) 조용히 빈 문자열을 돌려주지 않고
   `NotDocxError` 로 실제 정체를 담아 실패한다.

⚠️ 알려진 한계 — 세로 병합 셀은 원본 저장 도구에 따라 이어붙은 칸에도 같은 텍스트가
   중복 저장되어 있을 수 있다(가로 병합은 행 안에서 중복 제거한다). 텍스트박스/도형
   안 글상자는 `python-docx` 의 문서 순회 API 가 못 보므로 유실될 수 있다.
   둘 다 하위 우선순위 — 조문 재조립에 큰 영향이 없어 지금은 손대지 않는다.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from hwpx_extract import sniff, NotHwpxError  # noqa: F401  (sniff 재사용)


class NotDocxError(ValueError):
    """파일이 docx 가 아니다. 메시지에 실제 정체를 담는다."""


def _walk_blocks(blocks, lines: list[str]) -> None:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for block in blocks:
        if isinstance(block, Paragraph):
            t = block.text
            if t.strip():
                lines.extend(t.split("\n"))
        elif isinstance(block, Table):
            _walk_table(block, lines)


def _walk_table(tbl, lines: list[str]) -> None:
    """표 → 셀마다 한 줄. 가로 병합은 같은 tc 가 행 안에서 반복되므로 중복 제거한다."""
    for row in tbl.rows:
        seen: set[int] = set()
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            _walk_blocks(cell.iter_inner_content(), lines)


def extract(path: str | Path) -> str:
    """docx → 평문. 문단·표 셀 경계는 줄바꿈. docx 가 아니면 NotDocxError."""
    p = Path(path)
    kind = sniff(p)
    if not kind.startswith("DOCX"):
        raise NotDocxError(f"docx 아님 — 실제 정체: {kind} ({p.name})")

    from docx import Document

    doc = Document(str(p))
    lines: list[str] = []
    _walk_blocks(doc.iter_inner_content(), lines)
    return "\n".join(lines)


if __name__ == "__main__":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    text = extract(sys.argv[1])
    if len(sys.argv) > 2:
        open(sys.argv[2], "w", encoding="utf-8").write(text)
        print(f"wrote {len(text)} chars -> {sys.argv[2]}")
    else:
        print(text)
